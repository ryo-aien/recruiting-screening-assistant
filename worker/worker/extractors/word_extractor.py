import io
import logging

from docx import Document

logger = logging.getLogger(__name__)


class WordExtractor:
    """Extract text from Word documents (.docx)."""

    async def extract(self, content: bytes) -> str:
        """Extract text from Word document content.

        Args:
            content: Word document content as bytes

        Returns:
            Extracted text

        Raises:
            Exception: If extraction fails
        """
        try:
            doc_file = io.BytesIO(content)
            document = Document(doc_file)

            text_parts = []

            # Extract paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in document.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))

            if not text_parts:
                raise ValueError("No text could be extracted from Word document")

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            raise


def get_word_extractor() -> WordExtractor:
    return WordExtractor()
